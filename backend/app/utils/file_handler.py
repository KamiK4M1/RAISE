import os
import uuid
import aiofiles
from typing import List, Optional, BinaryIO
from pathlib import Path
import mimetypes
import logging
from PyPDF2 import PdfReader
from docx import Document
import asyncio
from llama_parse import LlamaParse

from app.config import settings
from app.core.exceptions import FileUploadError
from app.utils.thai_processing import thai_processor

logger = logging.getLogger(__name__)

class FileHandler:
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.max_file_size = settings.max_file_size
        self.allowed_types = settings.allowed_file_types
        
        # Create upload directory if it doesn't exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize LlamaParse
        self.llama_parser = None
        if settings.llamaparse_api_key:
            try:
                self.llama_parser = LlamaParse(
                    api_key=settings.llamaparse_api_key,
                    result_type="markdown",  # Can be "markdown" or "text"
                    num_workers=4,  # Number of parallel workers
                    verbose=True,
                )
                logger.info("LlamaParse initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize LlamaParse: {e}. Falling back to PyPDF2.")
                self.llama_parser = None
        else:
            logger.warning("LLAMAPARSE_API_KEY not found. Using PyPDF2 for PDF parsing.")

    def validate_file(self, filename: str, file_size: int) -> None:
        """Validate uploaded file"""
        # Check file size
        if file_size > self.max_file_size:
            raise FileUploadError(f"ไฟล์มีขนาดใหญ่เกินไป (สูงสุด {self.max_file_size // (1024*1024)} MB)")
        
        # Check file extension
        file_extension = filename.lower().split('.')[-1]
        if file_extension not in self.allowed_types:
            raise FileUploadError(f"ประเภทไฟล์ไม่รองรับ (รองรับเฉพาะ {', '.join(self.allowed_types)})")
        
        # Check filename
        if not filename or len(filename) > 255:
            raise FileUploadError("ชื่อไฟล์ไม่ถูกต้อง")

    async def save_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return file path"""
        try:
            # Generate unique filename
            file_extension = filename.split('.')[-1]
            unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
            file_path = self.upload_dir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            logger.info(f"File saved: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            raise FileUploadError(f"เกิดข้อผิดพลาดในการบันทึกไฟล์: {str(e)}")

    async def read_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using LlamaParse or PyPDF2 as fallback"""
        try:
            # Try LlamaParse first if available
            if self.llama_parser:
                try:
                    logger.info(f"Using LlamaParse to process PDF: {file_path}")
                    
                    # LlamaParse requires the file path directly
                    documents = await asyncio.to_thread(
                        self.llama_parser.load_data, file_path
                    )
                    
                    if documents and len(documents) > 0:
                        # Combine all document pages
                        text = ""
                        for doc in documents:
                            if hasattr(doc, 'text'):
                                text += doc.text + "\n"
                            elif hasattr(doc, 'content'):
                                text += doc.content + "\n"
                        
                        if text.strip():
                            logger.info("Successfully extracted text using LlamaParse")
                            return thai_processor.clean_thai_text(text)
                    
                    logger.warning("LlamaParse returned empty content, falling back to PyPDF2")
                    
                except Exception as e:
                    logger.warning(f"LlamaParse failed: {e}. Falling back to PyPDF2")
            
            # Fallback to PyPDF2
            logger.info(f"Using PyPDF2 to process PDF: {file_path}")
            
            def extract_pdf_text():
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    text = ""
                    
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    
                    return text

            # Run in thread pool to avoid blocking
            text = await asyncio.to_thread(extract_pdf_text)
            
            if not text.strip():
                raise FileUploadError("ไม่สามารถดึงข้อความจากไฟล์ PDF ได้")
            
            logger.info("Successfully extracted text using PyPDF2")
            return thai_processor.clean_thai_text(text)
            
        except FileUploadError:
            raise
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise FileUploadError(f"เกิดข้อผิดพลาดในการอ่านไฟล์ PDF: {str(e)}")

    async def read_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            def extract_docx_text():
                doc = Document(file_path)
                text = ""
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
                
                return text

            # Run in thread pool to avoid blocking
            text = await asyncio.to_thread(extract_docx_text)
            
            if not text.strip():
                raise FileUploadError("ไม่สามารถดึงข้อความจากไฟล์ Word ได้")
            
            return thai_processor.clean_thai_text(text)
            
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise FileUploadError(f"เกิดข้อผิดพลาดในการอ่านไฟล์ Word: {str(e)}")

    async def read_txt(self, file_path: str) -> str:
        """Read text from TXT file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
            
            if not text.strip():
                # Try different encodings
                encodings = ['utf-8', 'tis-620', 'windows-1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                            text = await f.read()
                        if text.strip():
                            break
                    except:
                        continue
            
            if not text.strip():
                raise FileUploadError("ไม่สามารถอ่านเนื้อหาจากไฟล์ข้อความได้")
            
            return thai_processor.clean_thai_text(text)
            
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise FileUploadError(f"เกิดข้อผิดพลาดในการอ่านไฟล์ข้อความ: {str(e)}")

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from file based on type"""
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return await self.read_pdf(file_path)
        elif file_type == 'docx':
            return await self.read_docx(file_path)
        elif file_type == 'txt':
            return await self.read_txt(file_path)
        else:
            raise FileUploadError(f"ประเภทไฟล์ {file_type} ไม่รองรับ")

    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[dict]:
        """Split text into chunks with overlap"""
        try:
            if not text.strip():
                return []
            
            # Split by sentences first
            sentences = thai_processor.segment_sentences(text)
            
            chunks = []
            current_chunk = ""
            current_size = 0
            chunk_index = 0
            
            for sentence in sentences:
                sentence_size = len(sentence)
                
                # If adding this sentence exceeds chunk size, start new chunk
                if current_size + sentence_size > chunk_size and current_chunk:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'chunk_index': chunk_index,
                        'embedding': []  # Will be filled later
                    })
                    
                    # Start new chunk with overlap
                    overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                    current_chunk = overlap_text + " " + sentence
                    current_size = len(current_chunk)
                    chunk_index += 1
                else:
                    current_chunk += " " + sentence if current_chunk else sentence
                    current_size += sentence_size
            
            # Add final chunk if there's remaining text
            if current_chunk.strip():
                chunks.append({
                    'text': current_chunk.strip(),
                    'chunk_index': chunk_index,
                    'embedding': []
                })
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            return [{'text': text, 'chunk_index': 0, 'embedding': []}]

    async def delete_file(self, file_path: str) -> bool:
        """Delete uploaded file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File deleted: {file_path}")
                return True
            return False
            
        except Exception as e:
            logger.error(f"Error deleting file: {e}")
            return False

    def get_file_info(self, file_path: str) -> dict:
        """Get file information"""
        try:
            file_stat = os.stat(file_path)
            return {
                'size': file_stat.st_size,
                'created': file_stat.st_ctime,
                'modified': file_stat.st_mtime,
                'mime_type': mimetypes.guess_type(file_path)[0]
            }
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {}

# Global instance
file_handler = FileHandler()